import streamlit as st
import os
import io
import anthropic
from typing import List, Tuple, Dict
from dotenv import load_dotenv
import pdfplumber
import docx
import base64
import requests
import zipfile
from pdf2image import convert_from_bytes

load_dotenv()

# Configuration
CLAUDE_API_KEY = os.getenv("CLAUDE_API_KEY")
GOOGLE_VISION_API_KEY = os.getenv("GOOGLE_VISION_API_KEY")

class RAGPipeline:
    def __init__(self, google_vision_api_key):
        self.google_vision_api_key = google_vision_api_key

    def image_to_base64(self, image_path: str) -> str:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')

    def analyze_images(self, images: List[Tuple[str, Tuple[str, bytes, str]]]) -> str:
        texts = []
        for _, image_tuple in images[:15]:  # Limit to 15 pages for OCR
            try:
                _, img_bytes, _ = image_tuple
                base64_image = base64.b64encode(img_bytes).decode('utf-8')
                result = self.analyze_single_image(base64_image)
                text = result['responses'][0]['textAnnotations'][0]['description']
                texts.append(text)
            except Exception as e:
                st.error(f"Error processing image: {e}")
        return "\n\n".join(texts)

    def analyze_single_image(self, base64_image: str) -> Dict:
        url = f'https://vision.googleapis.com/v1/images:annotate?key={self.google_vision_api_key}'
        payload = {
            "requests": [
                {
                    "image": {
                        "content": base64_image
                    },
                    "features": [
                        {
                            "type": "TEXT_DETECTION",
                            "maxResults": 10
                        }
                    ]
                }
            ]
        }
        response = requests.post(url, json=payload)
        if response.status_code != 200:
            raise Exception(f"Request failed with status code {response.status_code}: {response.text}")
        return response.json()

def extract_text_from_file(file, rag_pipeline: RAGPipeline):
    file_extension = os.path.splitext(file.name)[1].lower()
    
    if file_extension == '.pdf':
        try:
            pdf_content = io.BytesIO(file.read())
            file.seek(0)  # Reset file pointer
            
            with pdfplumber.open(pdf_content) as pdf:
                first_page = pdf.pages[0]
                text = first_page.extract_text() or ""
                if len(text) > 100:
                    return " ".join([page.extract_text() or "" for page in pdf.pages])
                else:
                    images = convert_from_bytes(file.read())
                    image_files = []
                    for i, image in enumerate(images[:25]):  # Limit to 25 pages for OCR
                        img_byte_arr = io.BytesIO()
                        image.save(img_byte_arr, format='JPEG')
                        img_byte_arr = img_byte_arr.getvalue()
                        image_files.append(('image', ('image.jpg', img_byte_arr, 'image/jpeg')))
                    
                    ocr_text = rag_pipeline.analyze_images(image_files)
                    return ocr_text
        except Exception as e:
            st.error(f"Error processing PDF: {e}")
            return ""
    elif file_extension in ['.doc', '.docx']:
        doc = docx.Document(io.BytesIO(file.read()))
        return " ".join([paragraph.text for paragraph in doc.paragraphs])
    elif file_extension in ['.txt']:
        return file.read().decode('utf-8')
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

def extract_text_from_zip(zip_file, rag_pipeline: RAGPipeline):
    extracted_texts = {}
    with zipfile.ZipFile(zip_file, 'r') as zip_ref:
        for file_name in zip_ref.namelist():
            try:
                with zip_ref.open(file_name) as file:
                    file_like_object = io.BytesIO(file.read())
                    file_like_object.name = file_name  # Add name attribute
                    text = extract_text_from_file(file_like_object, rag_pipeline)
                    if text.strip():
                        extracted_texts[file_name] = text
                        st.success(f"Successfully extracted text from {file_name}")
                    else:
                        st.warning(f"No text could be extracted from {file_name}")
                        extracted_texts[file_name] = "No text could be extracted"
            except Exception as e:
                st.error(f"Error processing {file_name} from ZIP: {str(e)}")
                extracted_texts[file_name] = f"Error: {str(e)}"
    return extracted_texts

def claudeCall(text, prompt):
    client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)

    try:
        message = client.messages.create(
            model="claude-3-5-sonnet-20240620",
            max_tokens=3000,
            temperature=0,
            system="You are a general counsel of a fortune 500 company.",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": text,
                        },
                        {
                            "type": "text",
                            "text": prompt
                        },
                    ],
                }
            ],
        )
    except Exception as e:
        st.error(f"An error occurred while calling Claude API: {e}")
        return ""

    text_content = ''.join(block.text for block in message.content if block.type == 'text')
    return text_content

def perform_conflict_check(extracted_texts: Dict[str, str]) -> str:
    if len(extracted_texts) < 2:
        return "At least two documents are required for a conflict check."

    combined_text = ""
    for filename, text in extracted_texts.items():
        combined_text += f"Document: {filename}\n\n{text}\n\n---\n\n"

    conflict_check_prompt = """
    Perform a conflict check across all the provided documents. For each document, identify any clauses or terms that may conflict with clauses or terms in the other documents. 

    Provide your analysis in the following format:

    Document: [Filename1]
    Conflicts:
    1. Clause [X] conflicts with [Filename2], Clause [Y]:
       - [Brief explanation of the conflict]
    2. ...

    Document: [Filename2]
    Conflicts:
    1. ...

    If no conflicts are found for a document, state "No conflicts found."

    Focus on significant conflicts that could impact the legal or business relationship between the parties involved.
    """

    return claudeCall(combined_text, conflict_check_prompt)

def main():
    st.set_page_config(page_title="Order! Order!", layout="wide")
    
    st.title("Order! Order!")
    st.markdown("Upload your legal documents for quick analysis and summary. Supports PDF, DOCX, DOC, TXT, and ZIP files containing these formats.")

    st.success("All systems online!")

    # Initialize RAGPipeline
    rag_pipeline = RAGPipeline(GOOGLE_VISION_API_KEY)

    # Initialize session state variables
    if "extracted_texts" not in st.session_state:
        st.session_state.extracted_texts = {}
    if "messages" not in st.session_state:
        st.session_state.messages = []
    if "current_prompt" not in st.session_state:
        st.session_state.current_prompt = ""
    if "summary_expanded" not in st.session_state:
        st.session_state.summary_expanded = False
    if "risky_analysis_expanded" not in st.session_state:
        st.session_state.risky_analysis_expanded = False
    if "conflict_check_expanded" not in st.session_state:
        st.session_state.conflict_check_expanded = False
    if "ask_assistant_expanded" not in st.session_state:
        st.session_state.ask_assistant_expanded = False

    # Top options for single/multiple contracts
    analysis_type = st.radio("Select analysis type:", ["Single Contract", "Multiple Contracts"], horizontal=True)

    # File uploader
    if analysis_type == "Single Contract":
        uploaded_file = st.file_uploader("Upload document", type=['pdf', 'docx', 'doc', 'txt'])
        if uploaded_file:
            uploaded_files = [uploaded_file]
        else:
            uploaded_files = []
    else:
        uploaded_files = st.file_uploader("Upload document(s)", accept_multiple_files=True, type=['pdf', 'docx', 'doc', 'txt', 'zip'])

    # Clear previous data when new files are uploaded
    if uploaded_files:
        st.session_state.extracted_texts = {}
        st.session_state.messages = []
        st.session_state.summary_expanded = False
        st.session_state.risky_analysis_expanded = False
        st.session_state.conflict_check_expanded = False

        # Display uploaded documents and extract text
        st.subheader("Uploaded Document(s)")
        for file in uploaded_files:
            with st.expander(f"View {file.name}"):
                if file.name.lower().endswith('.zip'):
                    if analysis_type == "Multiple Contracts":
                        extracted_texts = extract_text_from_zip(file, rag_pipeline)
                        for filename, text in extracted_texts.items():
                            st.subheader(f"Content of {filename}")
                            st.text_area(f"Document Content - {filename}", text, height=300)
                            st.session_state.extracted_texts[filename] = text
                    else:
                        st.warning("ZIP files are not allowed for Single Contract analysis.")
                else:
                    text = extract_text_from_file(file, rag_pipeline)
                    if text:
                        st.text_area("Document Content", text, height=300)
                        st.session_state.extracted_texts[file.name] = text
                    else:
                        st.warning(f"Could not extract text from {file.name}. Please check if the file is valid and in a supported format.")

    # Sidebar
    st.sidebar.title("Analysis Options")

    # Summary Section
    if st.sidebar.button("Summary"):
        st.session_state.summary_expanded = not st.session_state.summary_expanded
    
    if st.session_state.summary_expanded:
        with st.sidebar.expander("Summary", expanded=True):
            if st.session_state.extracted_texts:
                for filename, text in st.session_state.extracted_texts.items():
                    if text:
                        st.write(f"Generating summary for {filename}...")
                        summary = claudeCall(text, "Provide a brief summary of this document.")
                        st.subheader(f"Summary of {filename}")
                        st.write(summary)
                        st.write("---")
                    else:
                        st.warning(f"Could not extract text from {filename}. Please check if the file is corrupted or in an unsupported format.")
            else:
                st.warning("Please upload a document first.")

    # Risky Analysis Section
    if st.sidebar.button("Risky Analysis"):
        st.session_state.risky_analysis_expanded = not st.session_state.risky_analysis_expanded

    if st.session_state.risky_analysis_expanded:
        with st.sidebar.expander("Risky Analysis", expanded=True):
            if st.session_state.extracted_texts:
                for filename, text in st.session_state.extracted_texts.items():
                    if text:
                        st.write(f"Performing risky analysis for {filename}...")
                        analysis_prompt = """
                        Analyze the document and identify potentially risky clauses or terms. For each risky clause:
                        1. Start with the actual clause number as it appears in the document.
                        2. Quote the relevant part of the clause.
                        3. Explain why it's potentially risky.

                        Format your response as follows:

                        Clause [X]: "[Quote the relevant part]"
                        Risk: [Explain the potential risk]

                        Where [X] is the actual clause number from the document.
                        IF NO CLAUSE NUMBER IS PRESENT IN THE DOCUMENT, DO NOT GIVE ANY NUMBER TO THE CLAUSE BY YOURSELF THEN.
                        """
                        analysis = claudeCall(text, analysis_prompt)
                        st.subheader(f"Risky Analysis of {filename}")
                        st.write(analysis)
                        st.write("---")
                    else:
                        st.warning(f"Could not extract text from {filename}. Please check if the file is corrupted or in an unsupported format.")
            else:
                st.warning("Please upload a document first.")

    # Conflict Check Section (Only for Multiple Contracts)
    if analysis_type == "Multiple Contracts":
        if st.sidebar.button("Conflict Check"):
            st.session_state.conflict_check_expanded = not st.session_state.conflict_check_expanded

        if st.session_state.conflict_check_expanded:
            with st.sidebar.expander("Conflict Check", expanded=True):
                if len(st.session_state.extracted_texts) >= 2:
                    st.write("Performing conflict check across all documents...")
                    conflict_analysis = perform_conflict_check(st.session_state.extracted_texts)
                    st.subheader("Conflict Check Results")
                    st.write(conflict_analysis)
                elif len(st.session_state.extracted_texts) == 1:
                    st.warning("At least two documents are required for a conflict check. Please upload more documents.")
                else:
                    st.warning("Please upload at least two documents for a conflict check.")

    # Ask Assistant Section
    if st.sidebar.button("Ask Assistant"):
        st.session_state.ask_assistant_expanded = not st.session_state.ask_assistant_expanded
    
    if st.session_state.ask_assistant_expanded:
        with st.sidebar.expander("Ask Assistant", expanded=True):
            st.write("Chat with the AI Assistant about the uploaded document(s).")
            
            # Display chat history
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

            # Chat input
            prompt = st.text_input("What would you like to know about the document(s)?", key="chat_input", value=st.session_state.current_prompt)
            send_button = st.button("Send")

            # Check if Enter was pressed (prompt changed) or Send button was clicked
            if prompt != st.session_state.current_prompt or send_button:
                if prompt.strip():  # Ensure the prompt is not empty
                    st.session_state.current_prompt = prompt  # Update the current prompt
                    st.session_state.messages.append({"role": "user", "content": prompt})
                    
                    with st.chat_message("user"):
                        st.markdown(prompt)

                    with st.chat_message("assistant"):
                        if st.session_state.extracted_texts:
                            full_text = "\n\n".join(st.session_state.extracted_texts.values())
                            
                            if full_text.strip():
                                response = claudeCall(full_text, prompt)
                                st.markdown(response)
                                st.session_state.messages.append({"role": "assistant", "content": response})
                            else:
                                st.warning("Could not extract text from the uploaded document(s). Please ensure the files are in a supported format and contain readable text.")
                                st.session_state.messages.append({"role": "assistant", "content": "Could not extract text from the uploaded document(s). Please ensure the files are in a supported format and contain readable text."})
                        else:
                            st.warning("Please upload a document first.")
                            st.session_state.messages.append({"role": "assistant", "content": "Please upload a document first."})
                    
                    # Clear the input after sending
                    st.session_state.current_prompt = ""

if __name__ == "__main__":
    main()
