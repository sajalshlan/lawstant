import os
import pdfplumber
import docx
from typing import Dict, Any, List
from langchain_community.llms import Predibase
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.text_splitter import RecursiveCharacterTextSplitter
import torch
import time
from langchain.embeddings.base import Embeddings
from transformers import AutoTokenizer, AutoModel
from dotenv import load_dotenv
import base64
import requests
from pdf2image import convert_from_path
import datetime

load_dotenv()

class InLegalBERTEmbeddings(Embeddings):
    def __init__(self, model_name="law-ai/InLegalBERT"):
        self.tokenizer = AutoTokenizer.from_pretrained(model_name)
        self.model = AutoModel.from_pretrained(model_name)
        self.device = 'cuda' if torch.cuda.is_available() else 'cpu'
        self.model.to(self.device)

    def embed_documents(self, texts):
        inputs = self.tokenizer(texts, padding=True, truncation=True, return_tensors="pt", max_length=512)
        inputs = {name: tensor.to(self.device) for name, tensor in inputs.items()}
        with torch.no_grad():
            outputs = self.model(**inputs)
        embeddings = outputs.last_hidden_state[:, 0, :].cpu().numpy()
        return embeddings.tolist()

    def embed_query(self, text):
        return self.embed_documents([text])[0]

class RAGPipeline:
    def __init__(self, embedding_model_name: str, llm_model_name: str, api_token: str, google_vision_api_key: str):
        print("Initializing RAGPipeline...")
        
        print(f"Loading embedding model: {embedding_model_name}")
        self.embed_model = self._create_embeddings(embedding_model_name)
        print("Embedding model loaded successfully.")
        
        print(f"Initializing Predibase model: {llm_model_name}")
        self.llm_model = self._create_predibase_model(llm_model_name, api_token)
        print("Predibase model initialized successfully.")
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        print("Text splitter initialized.")
        
        self.google_vision_api_key = google_vision_api_key
        print("Google Vision API key loaded.")
        
        print("RAGPipeline initialization complete.")

    @staticmethod
    def _create_embeddings(model_name: str) -> InLegalBERTEmbeddings:
        return InLegalBERTEmbeddings(model_name)

    @staticmethod
    def _create_predibase_model(model_name: str, api_token: str) -> Predibase:
        return Predibase(
            model=model_name,
            predibase_api_key=api_token,
            predibase_sdk_version=None,
            adapter_version=1,
            max_new_tokens=4000,
        )

    def save_file_with_timestamp(self, file, upload_folder: str) -> str:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename, file_extension = os.path.splitext(file.name)
        new_filename = f"{filename}_{timestamp}{file_extension}"
        file_path = os.path.join(upload_folder, new_filename)
        
        with open(file_path, "wb") as f:
            f.write(file.getvalue())
        
        return file_path

    def image_to_base64(self, image_path: str) -> str:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')

    def analyze_single_image(self, image_path: str) -> Dict:
        base64_image = self.image_to_base64(image_path)
        url = f'https://vision.googleapis.com/v1/images:annotate?key={self.google_vision_api_key}'
        payload = {
            "requests": [
                {
                    "image": {
                        "content": base64_image
                    },
                    "features": [
                        {
                            "type": "TEXT_DETECTION",
                            "maxResults": 10
                        }
                    ]
                }
            ]
        }
        response = requests.post(url, json=payload)
        if response.status_code != 200:
            raise Exception(f"Request failed with status code {response.status_code}: {response.text}")
        return response.json()

    def analyze_images(self, images: List[str]) -> str:
        texts = []
        for image_path in images[:15]:  # Limit to 15 pages for OCR
            try:
                result = self.analyze_single_image(image_path)
                text = result['responses'][0]['textAnnotations'][0]['description']
                texts.append(text)
            except Exception as e:
                print(f"Error processing image: {e}")
        return "\n\n".join(texts)

    # def extract_text_from_file(self, file_path: str) -> str:
    #     _, file_extension = os.path.splitext(file_path)
        
    #     if file_extension.lower() == '.pdf':
    #         # Check if the PDF is digitally native
    #         with pdfplumber.open(file_path) as pdf:
    #             first_page = pdf.pages[0]
    #             print('IN THE EXTRACT TEXT SECTION')
    #             text = first_page.extract_text() or ""
    #             print(len(text))
    #             if len(text) > 100:
    #                 # PDF is digitally native, use pdfplumber
    #                 print('DIGITALLY NATIVE')
    #                 print(len(text))
    #                 return " ".join([page.extract_text() or "" for page in pdf.pages])
    #             else:
    #                 # PDF needs OCR
    #                 print("*"*50, file_path)
    #                 images = convert_from_path(file_path)
    #                 image_paths = []
    #                 for i, image in enumerate(images[:10]):  # Limit to 10 pages for OCR
    #                     image_path = f"{file_path}_page_{i}.jpg"
    #                     image.save(image_path, "JPEG")
    #                     image_paths.append(image_path)
    #                 ocr_text = self.analyze_images(image_paths)
    #                 print(ocr_text)
    #                 # Clean up temporary image files
    #                 for image_path in image_paths:
    #                     os.remove(image_path)
    #                 return ocr_text
    #     elif file_extension.lower() in ['.doc', '.docx']:
    #         doc = docx.Document(file_path)
    #         return " ".join([paragraph.text for paragraph in doc.paragraphs])
    #     elif file_extension.lower() in ['.txt']:
    #         with open(file_path, 'r', encoding='utf-8') as file:
    #             return file.read()
    #     else:
    #         raise ValueError(f"Unsupported file format: {file_extension}")

    def extract_text_from_file(self, file_path: str) -> str:
        _, file_extension = os.path.splitext(file_path)
        
        if file_extension.lower() == '.pdf':
            try:
                # Check if the PDF is digitally native
                with pdfplumber.open(file_path) as pdf:
                    first_page = pdf.pages[0]
                    text = first_page.extract_text() or ""
                    print(len(text))
                    if len(text) > 100:
                        # PDF is digitally native, use pdfplumber
                        return " ".join([page.extract_text() or "" for page in pdf.pages])
                    else:
                        # PDF needs OCR
                        try:
                            images = convert_from_path(file_path)
                        except Exception as e:
                            print(f"Error converting PDF to images: {e}")
                            print("Make sure Poppler is installed and in your system PATH.")
                            return ""  # Return empty string if conversion fails
                        
                        image_paths = []
                        for i, image in enumerate(images[:25]):  # Limit to 25 pages for OCR
                            image_path = f"{file_path}_page_{i}.jpg"
                            image.save(image_path, "JPEG")
                            image_paths.append(image_path)
                        ocr_text = self.analyze_images(image_paths)
                        # Clean up temporary image files
                        for image_path in image_paths:
                            os.remove(image_path)
                        return ocr_text
            except Exception as e:
                print(f"Error processing PDF: {e}")
                return ""  # Return empty string if PDF processing fails
        elif file_extension.lower() in ['.doc', '.docx']:
            doc = docx.Document(file_path)
            return " ".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_extension.lower() in ['.txt']:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def create_vector_store(self, text: str) -> FAISS:
        documents = self.text_splitter.create_documents([text])
        return FAISS.from_documents(documents, embedding=self.embed_model)

    def create_rag_chain(self, vector_store: FAISS) -> RetrievalQA:
        retriever = vector_store.as_retriever()
        return RetrievalQA.from_chain_type(
            llm=self.llm_model,
            chain_type='stuff',
            retriever=retriever
        )

    @staticmethod
    def generate_prompt(system_message: str, query: str = "") -> str:
        return f"""
        <|begin_of_text|><|start_header_id|>system<|end_header_id|>
        {system_message}
        <|eot_id|>
        <|start_header_id|>user<|end_header_id|>
        Query: ```{query}```
        <|eot_id|><|start_header_id|>assistant<|end_header_id|>
        """

    def run_rag_pipeline(self, rag_chain: RetrievalQA, prompt: str) -> Dict[str, Any]:
        start_time = time.time()
        response = rag_chain.invoke(prompt)
        end_time = time.time()
        
        return {
            "result": response['result'],
            "execution_time": end_time - start_time
        }
